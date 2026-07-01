#!/usr/bin/env python3
"""
PRD docx 生成器

用法:
    python3 gen_prd_docx.py <config.json> [输出路径.docx]

config.json 结构见 prd_docx_schema_example.json

功能:
    - 根据 JSON 配置生成 PRD Word 文档
    - 在"原型/交互图"位置插入原型页面截图
    - 支持 heading/paragraph/table/quote/screenshot/page_break 等元素

依赖: python-docx
"""

import sys
import os
import json

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

DEFAULT_FONT = "Microsoft YaHei"


def set_font_run(run, name=DEFAULT_FONT, size=None, bold=False, color=None):
    run.font.name = name
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), name)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        set_font_run(run, size=None, bold=True)
    return h


def add_paragraph(doc, text, bold=False, color=None, size=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font_run(run, size=size or 11, bold=bold, color=color)
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                set_font_run(run, size=10, bold=True)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    set_font_run(run, size=10)
    return table


def add_screenshot(doc, path, caption, width_inches=5.5):
    if not os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("[原型截图: " + caption + "]")
        set_font_run(run, size=9, color=(0x99, 0x99, 0x99))
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Inches(width_inches))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    crun = cap.add_run(caption)
    set_font_run(crun, size=9, color=(0x99, 0x99, 0x99))


def add_quote(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    set_font_run(run, size=11, color=(0x66, 0x66, 0x66))
    return p


def build_cover(doc, config):
    cover = config.get("cover", {})
    title = cover.get("title", "未命名PRD")
    version = cover.get("version", "V1.0")
    date = cover.get("date", "")
    author = cover.get("author", "")

    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    set_font_run(run, size=24, bold=True)

    doc.add_paragraph()
    doc.add_paragraph()

    for item in [("版本: " + version), ("日期: " + date), ("作者: " + author)]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(item)
        set_font_run(run, size=12, color=(0x66, 0x66, 0x66))

    doc.add_page_break()


def render_content(doc, config, shot_dir=None):
    content = config.get("content", [])
    for item in content:
        t = item.get("type", "paragraph")
        if t == "heading":
            add_heading(doc, item.get("text", ""), level=item.get("level", 1))
        elif t == "paragraph":
            add_paragraph(doc, item.get("text", ""), bold=item.get("bold", False))
        elif t == "quote":
            add_quote(doc, item.get("text", ""))
        elif t == "table":
            add_table(doc, item.get("headers", []), item.get("rows", []))
        elif t == "screenshot":
            path = item.get("path", "")
            if shot_dir and item.get("file"):
                path = os.path.join(shot_dir, item.get("file"))
            add_screenshot(doc, path, item.get("caption", "原型截图"))
        elif t == "page_break":
            doc.add_page_break()


def main():
    if len(sys.argv) < 2:
        print("用法: python3 gen_prd_docx.py <config.json> [输出路径.docx]")
        sys.exit(1)

    config_path = sys.argv[1]
    output_path = sys.argv[2] if len(sys.argv) > 2 else None

    with open(config_path, "r", encoding="utf-8") as f:
        config = json.load(f)

    if not output_path:
        product_dir = config.get("output_dir", os.path.dirname(config_path))
        doc_filename = config.get("doc_filename", "PRD文档.docx")
        output_path = os.path.join(product_dir, doc_filename)

    shot_dir = config.get("screenshot_dir", None)

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = DEFAULT_FONT
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), DEFAULT_FONT)

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    build_cover(doc, config)
    render_content(doc, config, shot_dir)

    doc.save(output_path)
    print("Done: " + output_path)
    print("Size: " + str(round(os.path.getsize(output_path) / 1024, 1)) + " KB")


if __name__ == "__main__":
    main()
