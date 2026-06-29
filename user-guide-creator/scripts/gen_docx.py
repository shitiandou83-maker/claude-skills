#!/usr/bin/env python3
"""
用户手册 docx 生成器

用法:
    python3 gen_docx.py <config.json> [输出路径.docx]

config.json 结构见 doc_schema_example.json

依赖: python-docx
"""

import sys
import os
import json

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

# ------------------------------------------------------------
# 默认配置
# ------------------------------------------------------------
DEFAULT_FONT = "Microsoft YaHei"

# ------------------------------------------------------------
# 工具函数
# ------------------------------------------------------------
def set_font_run(run, name=DEFAULT_FONT, size=None, bold=False, color=None):
    """设置 run 的字体属性"""
    run.font.name = name
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    # 东亚字体回退
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), name)


def add_heading(doc, text, level=1):
    """添加标题，自动处理中文字体"""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        set_font_run(run, size=None, bold=True)
    return h


def add_paragraph(doc, text, bold=False, color=None, size=None, align=None):
    """添加段落"""
    p = doc.add_paragraph()
    if align:
        p.alignment = getattr(WD_ALIGN_PARAGRAPH, align.upper(), WD_ALIGN_PARAGRAPH.LEFT)
    run = p.add_run(text)
    set_font_run(run, size=size or 11, bold=bold, color=color)
    return p


def add_quote(doc, text, prefix=None):
    """添加引用块（缩进段落）"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    if prefix:
        run = p.add_run(prefix + "\n")
        set_font_run(run, size=11, bold=True)
    run = p.add_run(text)
    set_font_run(run, size=11)
    return p


def add_screenshot(doc, path, caption, width_inches=5.5):
    """插入图片 + 图注"""
    if not os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"[此处建议插入图片：{caption}]")
        set_font_run(run, size=9, color=(0x99, 0x99, 0x99))
        return

    # 图片
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Inches(width_inches))

    # 图注
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    crun = cap.add_run(caption)
    set_font_run(crun, size=9, color=(0x99, 0x99, 0x99))


def add_table(doc, headers, rows):
    """添加表格，自动格式化"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                set_font_run(run, size=10, bold=True)

    # 数据行
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    set_font_run(run, size=10)

    return table


def add_page_break(doc):
    """添加分页符"""
    doc.add_page_break()


def add_blank_lines(doc, count=1):
    """添加空行"""
    for _ in range(count):
        doc.add_paragraph()


def build_cover(doc, config):
    """生成封面"""
    cover = config.get("cover", {})
    title = cover.get("title", "未命名产品")
    subtitle = cover.get("subtitle", "使用指南")
    version = cover.get("version", "V1.0")
    date = cover.get("date", "")
    author = cover.get("author", "")

    for _ in range(6):
        doc.add_paragraph()

    # 标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    set_font_run(run, size=28, bold=True)

    # 副标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle)
    set_font_run(run, size=20)

    doc.add_paragraph()
    doc.add_paragraph()

    # 元信息
    meta_items = [("版本: " + version), ("日期: " + date), ("作者: " + author)]
    for item in meta_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(item)
        set_font_run(run, size=12, color=(0x66, 0x66, 0x66))

    add_page_break(doc)


# ------------------------------------------------------------
# 内容渲染
# ------------------------------------------------------------
def render_content(doc, config, shot_dir=None):
    """根据 JSON 配置渲染文档内容"""
    content = config.get("content", [])

    for item in content:
        t = item.get("type", "paragraph")

        if t == "heading":
            add_heading(doc, item.get("text", ""), level=item.get("level", 1))

        elif t == "paragraph":
            add_paragraph(doc, item.get("text", ""),
                        bold=item.get("bold", False),
                        color=item.get("color"),
                        size=item.get("size"),
                        align=item.get("align"))

        elif t == "quote":
            add_quote(doc, item.get("text", ""), prefix=item.get("prefix"))

        elif t == "list":
            bullets = item.get("bullets", []) if "bullets" in item else item.get("items", [])
            bold = item.get("bold", False)
            for b in bullets:
                if isinstance(b, dict):
                    add_paragraph(doc, "  - " + b.get("text", ""), bold=b.get("bold", bold))
                else:
                    add_paragraph(doc, "  - " + str(b), bold=bold)

        elif t == "table":
            add_table(doc, item.get("headers", []), item.get("rows", []))

        elif t == "screenshot":
            if shot_dir:
                path = os.path.join(shot_dir, item.get("file", ""))
            else:
                path = item.get("path", "")
            add_screenshot(doc, path, item.get("caption", "截图"))

        elif t == "page_break":
            add_page_break(doc)

        elif t == "blank":
            add_blank_lines(doc, item.get("count", 1))


# ------------------------------------------------------------
# 主入口
# ------------------------------------------------------------
def main():
    if len(sys.argv) < 2:
        print("用法: python3 gen_docx.py <config.json> [输出路径.docx]")
        print("config.json 格式见 doc_schema_example.json")
        sys.exit(1)

    config_path = sys.argv[1]
    output_path = sys.argv[2] if len(sys.argv) > 2 else None

    # 读取配置
    with open(config_path, "r", encoding="utf-8") as f:
        config = json.load(f)

    # 推断输出路径
    if not output_path:
        product_dir = config.get("output_dir", os.path.dirname(config_path))
        doc_filename = config.get("doc_filename", "用户手册.docx")
        output_path = os.path.join(product_dir, doc_filename)

    # 截图目录
    shot_dir = config.get("screenshot_dir", None)

    # 创建文档
    doc = Document()

    # 全局样式
    style = doc.styles['Normal']
    style.font.name = DEFAULT_FONT
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), DEFAULT_FONT)

    # 页边距
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # 封面
    build_cover(doc, config)

    # 渲染内容
    render_content(doc, config, shot_dir)

    # 保存
    doc.save(output_path)
    print(f"✅ 已生成: {output_path}")
    print(f"📏 大小: {os.path.getsize(output_path)/1024:.1f} KB")


if __name__ == "__main__":
    main()
